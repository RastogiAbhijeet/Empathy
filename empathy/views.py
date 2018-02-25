from django.shortcuts import render, redirect
from django.http import HttpResponse, HttpResponseRedirect
from django.urls import reverse
from django.views.decorators.csrf import csrf_exempt
import json 
import csv
from .models import LoginTable, StudentProfile, EventTable, RollNo
import mimetypes
from django.http import StreamingHttpResponse
from wsgiref.util import FileWrapper
import os
import time
import random
from docx import Document
from docx.shared import Pt

from django.core.exceptions import ObjectDoesNotExist
from django.core.mail import EmailMessage

# Create your views here. 

@csrf_exempt 
def validate(request):
    # Login Validation - 
    #     Checkes whether the received username and password are valid or not
    # responseText = ""
    js = json.loads(request.body.decode('utf-8'))
    userName = js['Username'] 
    passWord = js['Password']
    # r

    print("Username : "+userName+"Password : " + passWord)
    try:
        dbObj = LoginTable.objects.filter(username = userName)
        print(dbObj)
        try:
            dic = {}
            dbObj.get(password = str(passWord))
            for i in dbObj:
                print(i.role)

                if str(i.role) == "student":
                    studentObj = StudentProfile.objects.get(roll_no = passWord)
                    dic = {
                        'name':studentObj.name,
                        'roll_no':studentObj.roll_no,
                        'role':studentObj.role,
                        'gender':studentObj.gender,
                        'branch':studentObj.branch,
                        'CCET':studentObj.college,
                        'event_type':studentObj.event_type,
                        'event_year':studentObj.event_year
                    }
                else:
                    dic = {'role':str(i.role)}
            print(dic)
            ls = [dic]
            responseText = str(ls)[1:-1]

        except ObjectDoesNotExist:
            responseText = ''
            # return HttpResponse(responseText, content_type="text/plain")
        except Exception as e:
            print(e)

        return HttpResponse(responseText, content_type="text/plain")
    except:     
        return HttpResponse("No User Found", content_type="text/plain")

@csrf_exempt
def register(request):
    
    dbObj = StudentProfile()
    loginObj = LoginTable()
    js = json.loads(request.body.decode('utf-8'))

    try:
        StudentProfile.objects.get(roll_no = str(js["roll_no"]))
        print("Double Help")
        return HttpResponse("Record Exists", content_type="text/plain")

    except ObjectDoesNotExist:

        if checkValidRollNo((str(js["roll_no"]).upper())):
            print(ObjectDoesNotExist)
            print("Hell")
            dbRoll = RollNo.objects.filter(roll_no=str(js['roll_no']).upper())
            name = dbRoll[0].name
            dbObj.name = str(name)
            dbObj.roll_no = str(js["roll_no"])
            dbObj.branch = str(js["branch"])
            dbObj.semester = str(js["semester"])
            dbObj.year = str(js["year"])
            dbObj.email = str(js["email"])
            dbObj.mobile = str(js["mobile"])
            dbObj.role = str(js["role"])
            dbObj.gender = str(js["gender"])
            dbObj.college = str(js["college"])
            dbObj.event_year = time.strftime("%Y")
            dbObj.save()

            loginObj.password =str(js["roll_no"])
            loginObj.username = str(js["roll_no"])
            loginObj.role = str(js["role"])
            loginObj.save()
            
            return HttpResponse("success", content_type= "text/plain")
        else :
            return HttpResponse("Can't find Roll No in the records", content_type = "text/plain")

    except :
        return HttpResponse("Please Update Your App or Try After sometime", content_type = "text/plain")

def checkValidRollNo(roll):
    
    print("In Validation Shitt !! ")
    try:
        db = RollNo.objects.get(roll_no = roll)
        return True
    except:
        print("Hello")
        return False

# Here's How we are going to send the PDF.
@csrf_exempt
def output_file(request):
    file_name = request.body.decode('ascii')
    path_to_notices = os.path.join(os.getcwd(), 'media/notices')
    list_notices = os.listdir(path_to_notices)
    
    if file_name in list_notices:
        file_path = "media/notices/" + file_name;
        print(file_path)
    
        with open(file_path,"rb") as f:
            data = f.read()

        response = HttpResponse(data, content_type = mimetypes.guess_type(url = file_path)[0])
        response['Content-Disposition'] = "attachment;filename = 'data'"
        response['Content-Length'] = os.path.getsize(file_path)
        response['File-Name'] = file_name

        return response
    
    else:
        return HttpResponse("", content_type = "text/plain")

@csrf_exempt
def event_register(request):
    
    try:
        # print("Hello")
        js = json.loads(request.body.decode("utf-8"))

        if js['event_type'] == "Inter Year":
            jsArray = js['event']
            dbObj = EventTable.objects.filter(roll_no = str(js["roll_no"]))
            
            if(len(dbObj) < 8):
                
                for events in jsArray:
                    db = EventTable()
                    db.roll_no = str(js["roll_no"])
                    db.event = events
                    db.position = "Not Played"
                    db.event_type = str(js['event_type'])
                    db.save()
                
            else: 
                response_str = "Already have registered for events"
                for i in dbObj:
                    response_str = response_str + "   " +str(i.event)

            return HttpResponse(response_str, content_type="text/plain")

        elif js['event_type'] == "Athletic Meet":
            jsArray = js['event']
            dbObj = EventTable.objects.filter(roll_no = str(js["roll_no"]))
            list_registered = []
            for x in dbObj:
                list_registered.append(x)
                
            
            response_str = ""
            if(len(dbObj) <= 5):
                for i in jsArray:
                    if i not in list_registered:
                        db = EventTable()
                        print(i)
                        db.roll_no = js["roll_no"]
                        db.event = i
                        db.position = "Not Played"
                        db.save()
                response_str = "Check profile to see the registered Events"
            else: 
                response_str = "Already have registered for events"
                for i in dbObj:
                    response_str = response_str + "   " +str(i.event)

            return HttpResponse(response_str, content_type="text/plain")

    except Exception as e: 
        print(str(e))
        return HttpResponse(str(e), content_type="text/plain")
    # Js Data : {"name":"abhijeet rastogi","roll_no", Event:["event1","event2, "event3", "event4", "event5"]}

@csrf_exempt
def event_push(request):
    
    try:
        db = EventTable.objects.filter(roll_no=request.body.decode("utf-8"))

        event_list = [x.event for x in db]
        position_list = [x.position for x in db]
        event_type = [x.event_type for x in db]

        event_data = []
        for i in range(0, len(event_list)):
            event_data.append(str(event_list[i])+" | "+str(event_type[i] + " -- " + str(position_list[i])))


        print(position_list)
        print(event_type)


        dic = {"event":event_data}
        ls = [dic]
        responseText = str(ls)[1:-1]

        return HttpResponse(responseText, content_type = "text/plain")
    except Exception as e:
        print(e)

@csrf_exempt
def sendnames(request):
    
    js = json.loads(request.body.decode("utf-8"))
    db = EventTable.objects.filter(event = js["event"])
    
    responseArray = []

    for i in db:
        studentDic = {}
    
        dbObj = StudentProfile.objects.filter(roll_no = str(i.roll_no))
        if dbObj[0].gender == js["gender"]:
            studentDic['roll_no'] = str(dbObj[0].roll_no)
            studentDic['name'] = str(dbObj[0].name)
            studentDic['year'] = str(dbObj[0].year)
            studentDic['branch'] = str(dbObj[0].branch)
            studentDic['position'] = i.position
            studentDic['college'] = str(dbObj[0].college)
            responseArray.append(studentDic)

    dic = {"student":responseArray}
    ls = [dic]
    dataPacket = str(ls)[1:-1]
    print(dataPacket)

    return HttpResponse(dataPacket, content_type = "text/plain")
        
@csrf_exempt
def updateShortList(request):
    requestData = json.loads((request.body.decode('utf-8')))
    for dataResource in requestData:
        EventTable.objects.filter(event = dataResource['event'], roll_no = dataResource['roll_no']).update(position = dataResource['position'])
    return HttpResponse("Success", content_type = "text/plain")

@csrf_exempt
def updatePositionList(request):
    requestData = json.loads((request.body.decode('utf-8')))
    for dataResource in requestData:
        EventTable.objects.filter(event = dataResource['event'], roll_no = dataResource['roll_no'], position = dataResource['qualified']).update(position = dataResource['position'])
    return HttpResponse("Success", content_type = "text/plain")

@csrf_exempt
def sendMail(request):
    
    email = request.body.decode("utf-8")
    print(email)
    pin = random.randint(100000, 999999)
    obj = EmailMessage("Athletic Meet Registeration",str(pin), to=[email])
    obj.send()

    

    return HttpResponse(str(pin), content_type = "text/plain")

@csrf_exempt
def reportGeneration(request):
    
    gender_mask = json.loads(request.body.decode('utf-8'))
    roll_no = []

    for gender in gender_mask['gender']:
        db = StudentProfile.objects.filter(gender = gender)
        for dbInstance in db:
            roll_no.append(dbInstance.roll_no)
        
    finalRolls = []
    for roll in roll_no:
        dbEvent = EventTable.objects.filter(roll_no = roll)
        for event in dbEvent:
            if str(event.position) in gender_mask['position']:
                db = StudentProfile.objects.filter(roll_no= roll)
                dic = {"name": str(db[0].name),
                    "roll_no":str(db[0].roll_no), 
                    "branch":str(db[0].branch),
                    "year":str(db[0].year),
                    "event":str(event.event),
                    "position":str(event.position)
                }
                finalRolls.append(dic)

    employ_data = open('./media/EmployData.csv', 'a')

    # create the csv writer object

    csvwriter = csv.writer(employ_data)

    count = 0
    for x in finalRolls:
        js= json.dumps(x)
        jsoe = json.loads(js)
        print(jsoe.keys()) 
        if(count == 0):
            header = jsoe.keys()
            csvwriter.writerow(header)
            count+=1
        csvwriter.writerow(jsoe.values())

    employ_data.close()

    return HttpResponse(str(finalRolls)[1:-1],content_type="text/plain")


def loadCsv(request):
    
    file = open("./media/CollegeCutList.csv", "r")
    reader = csv.reader(file)
    for i in reader:
        db = RollNo()
        if(i[0] != 0):
            db.roll_no = i[0]
            db.name = i[1]
            db.fatherName = i[2]
        db.save()
    
    return HttpResponse("Data Successfully Loaded")

# /noticeList
@csrf_exempt
def sendNoticeList(request):
    print("Hello")
    try:

        path_to_notices = os.path.join(os.getcwd(), 'media/notices')
        list_notices = os.listdir(path_to_notices)
        result_dic = {"files":list_notices}
        result_dic = [result_dic]        
        result_dic = str(result_dic)[1:-1]

        print("Sending Data : " + result_dic)

        response = HttpResponse(result_dic, content_type = "text/plain")
        response['file_name'] = "Hello"
        return response
    
    except Exception as e:
        print(e)
        return HttpResponse("Hello", content_type = "text/plain")

@csrf_exempt
def certificate_push(request):
    js = json.loads(request.body.decode('utf-8'))
    
    try:
        os.mkdir("TempCertificate")
    except FileExistsError:
        pass
    
    certificate_dic = {}

    # db = StudentProfile.objects.filter
    certificate_dic["roll_no"] = str(js['roll_no'])
    certificate_dic["name"] = str(js['name'])
    certificate_dic["event"] = str(js['event'])
    certificate_dic["position"] = str(js['position'])
    certificate_dic["branch"] = str(js['branch'])

    
    father = RollNo.objects.filter(roll_no = str(js['roll_no']).upper())

    for i in father:
        print(i.fatherName)
        certificate_dic["fatherName"] = i.fatherName
    certificate_dic["position"] = str(js['position'])


    print(certificate_dic)
    generateDoc(certificate_dic)

    file_name = "certi" +".docx"
    file_path = "TempCertificate/" +file_name;
    
    print(file_path)

    with open(file_path,"rb") as f:
        data = f.read()

    response = HttpResponse(data, content_type = mimetypes.guess_type(url = file_path)[0])
    response['Content-Disposition'] = "attachment;filename = 'data'"
    response['Content-Length'] = os.path.getsize(file_path)
    response['File-Name'] = str(js['roll_no'])+str(js["event"]) + ".docx"

    return response


def generateDoc(jsObj):
    
    f = open("media/certificate/Atheletics Certificate Position 2017.docx",'rb')
    document = Document(f)
    style = document.styles

    for i in range(len(document.paragraphs)):
        
        if("&NAME&" in document.paragraphs[i].text):
            document.paragraphs[i].text = document.paragraphs[i].text.replace("&NAME&", jsObj['name'])
            

        if("&FATHER&" in document.paragraphs[i].text):
            document.paragraphs[i].text = document.paragraphs[i].text.replace("&FATHER&", jsObj['fatherName'])
            
        
        if("&ROLL&" in document.paragraphs[i].text):
            document.paragraphs[i].text = document.paragraphs[i].text.replace("&ROLL&", jsObj["roll_no"])

        if("&DEPT&" in document.paragraphs[i].text):
            
            if(jsObj["branch"] == "CSE"):
                branch = "Computer Science and Engineering"
            elif (jsObj["branch"] == "MECH") :
                branch = "Mechanical Engineering"
            elif (jsObj["branch"] == "ECE") :
                branch = "Electronics and Communication Engineering"
            elif (jsObj["branch"] == "CIVIL") :
                branch = "Civil Engineering"
            elif (jsObj["branch"] == "PROD") :
                branch = "Production and Industrial Engineering"
            elif (jsObj["branch"] == "ARCH") :
                branch = "Architecture"
            elif (jsObj["branch"] == "ELECT") :
                branch = "Electrical  Engineering"
           
                

            document.paragraphs[i].text = document.paragraphs[i].text.replace("&DEPT&",branch)
            
        
        if("&POSITION&" in document.paragraphs[i].text):
            document.paragraphs[i].text = document.paragraphs[i].text.replace("&POSITION&", jsObj["position"])

        if("&EVENT&" in document.paragraphs[i].text):
            document.paragraphs[i].text = document.paragraphs[i].text.replace("&EVENT&", jsObj["event"])
        

            # document.paragraphs[i].style = style['Body Text']
            document.paragraphs[i].style.font.size = Pt(20)
            document.paragraphs[i].style.font.name = 'Monotype Corsiva'
            document.paragraphs[i].style.font.bold = True

   
    file_name = "certi"+".docx"
    document.save("TempCertificate/"+file_name)

def bulkReportGenerationAthleticMeet(request):
    
    ls = []
    
    db = StudentProfile.objects.filter()

    for studentInstance in db:
        dic = {}
        dic["Name"] = studentInstance.name
        dic["Roll"] = studentInstance.roll_no
        dic["Semester"] = studentInstance.semester
        dic["Branch"] = studentInstance.branch
        dic["Mobile"] = studentInstance.mobile

        temp_event_list = []
        eventdb = EventTable.objects.filter(event_type = "Inter Year", roll_no = studentInstance.roll_no, event = "Chess")
        # print("Heelo")
        for eventInstance in eventdb:
            temp_event_list.append(eventInstance.event)
        
        if len(temp_event_list) != 0:
            dic["Event"] = temp_event_list
            ls.append(dic)
     
    
    hello_file = open("./media/Ehe.csv", 'w', newline="")
    csvwriter = csv.writer(hello_file)

    count = 0
    for x in ls:
        if count == 0:
            
            key = [v for v in x.keys()]
            header = key
            csvwriter.writerow(header)
            count+=1

        value = [v for v in x.values()]
        csvwriter.writerow(value)

    hello_file.close()

    with open("./media/Ehe.csv", "rb") as f:
        data = f.read()


    response = HttpResponse(data, content_type =  mimetypes.guess_type(url = "./media/Ehe.csv")[0])
    response['Content-Disposition'] = 'attachment; filename = "somefile.csv'

    return response   

@csrf_exempt
def deleteEvent(request):
    try:
        js = json.loads(request.body.decode("utf-8"))

        eventSplit = str(js['event']).split(' |')
    
        EventTable.objects.filter(roll_no = str(js['roll_no']), event = eventSplit[0]).delete()
        return HttpResponse("Success in deletion of event", content_type = "text/plain")
    except Exception as e:
        print(e)
        return HttpResponse("Failure to delete event", content_type = "text/plain")
    

    